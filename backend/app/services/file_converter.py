"""文件转换服务 — 将各种格式转为 Markdown"""

import csv
import io
import base64
from pathlib import Path
from typing import Optional

from app.services.agnes_ai import agnes_ai


class FileConverter:
    """文件到 Markdown 的转换器"""

    SUPPORTED_EXTENSIONS = {
        ".md": "markdown",
        ".txt": "text",
        ".csv": "csv",
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "excel",
        ".xls": "excel",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".webp": "image",
    }

    @classmethod
    def detect_type(cls, filename: str) -> Optional[str]:
        ext = Path(filename).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext)

    @classmethod
    async def convert(cls, file_bytes: bytes, filename: str) -> str:
        file_type = cls.detect_type(filename)
        if not file_type:
            raise ValueError(f"Unsupported file type: {filename}")

        converters = {
            "markdown": cls._convert_markdown,
            "text": cls._convert_text,
            "csv": cls._convert_csv,
            "pdf": cls._convert_pdf,
            "docx": cls._convert_docx,
            "excel": cls._convert_excel,
            "image": cls._convert_image,
        }
        converter = converters[file_type]
        # Sync converters don't need await, async ones do
        if file_type in ("markdown", "text", "csv"):
            return converter(file_bytes, filename)
        return await converter(file_bytes, filename)

    @staticmethod
    def _convert_markdown(data: bytes, filename: str) -> str:
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _convert_text(data: bytes, filename: str) -> str:
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _convert_csv(data: bytes, filename: str) -> str:
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return ""
        headers = rows[0]
        col_widths = [len(h) for h in headers]
        for row in rows[1:]:
            for i, cell in enumerate(row):
                col_widths[i] = max(col_widths[i], len(str(cell)))
        md_lines = ["| " + " | ".join(h.ljust(col_widths[i]) for i, h in enumerate(headers)) + " |"]
        md_lines.append("| " + " | ".join("-" * w for w in col_widths) + " |")
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(
                str(row[i]).ljust(col_widths[i]) if i < len(row) else "".ljust(col_widths[0])
                for i in range(len(headers))
            ) + " |")
        return "\n".join(md_lines)

    @staticmethod
    async def _convert_pdf(data: bytes, filename: str) -> str:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages) if pages else ""

    @staticmethod
    async def _convert_docx(data: bytes, filename: str) -> str:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n\n".join(parts)

    @staticmethod
    async def _convert_excel(data: bytes, filename: str) -> str:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheets = []
        for ws_name in wb.sheetnames:
            ws = wb[ws_name]
            rows_data = list(ws.values)
            if not rows_data:
                continue
            headers = rows_data[0]
            col_widths = [len(str(h)) for h in headers]
            for row in rows_data[1:]:
                for i, cell in enumerate(row):
                    if cell is not None:
                        col_widths[i] = max(col_widths[i], len(str(cell)))
            md_lines = ["## Sheet: " + ws_name, ""]
            md_lines.append("| " + " | ".join(str(h).ljust(col_widths[i]) for i, h in enumerate(headers)) + " |")
            md_lines.append("| " + " | ".join("-" * w for w in col_widths) + " |")
            for row in rows_data[1:]:
                md_lines.append("| " + " | ".join(
                    str(row[i]).ljust(col_widths[i]) if i < len(row) and row[i] is not None else "".ljust(col_widths[0])
                    for i in range(len(headers))
                ) + " |")
            sheets.append("\n".join(md_lines))
        wb.close()
        return "\n\n".join(sheets)

    @staticmethod
    async def _convert_image(data: bytes, filename: str) -> str:
        suffix = Path(filename).suffix.lstrip(".").lower()
        b64 = base64.b64encode(data).decode("utf-8")
        data_url = f"data:image/{suffix};base64,{b64}"
        result = await agnes_ai.image_recognition(
            data_url,
            "请详细描述这张图片中的所有文字和内容，保留原始格式"
        )
        return result["choices"][0]["message"]["content"]
